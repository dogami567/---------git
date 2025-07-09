"""
报告格式化服务

提供将Markdown格式报告转换为其他格式的功能。
"""

import os
import re
import json
import logging
from typing import Dict, Any, List, Optional, Union, Tuple, TYPE_CHECKING
from datetime import datetime
from pathlib import Path
from enum import Enum
from io import BytesIO

from backend.app.core.logging import logger
from backend.app.models.report import ReportTemplate, ReportSection, ReportFormatEnum

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("未安装python-docx库，DOCX功能将不可用")

try:
    import markdown
    from markdown.extensions import Extension
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    # 如果markdown库不可用，定义一个空的Extension类
    class Extension:
        def extendMarkdown(self, md):
            pass

try:
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name
    from pygments.formatters import HtmlFormatter
    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False

from backend.app.core.config import settings
from backend.app.core.exceptions import ReportGenerationException

# 使用TYPE_CHECKING避免循环导入
if TYPE_CHECKING:
    from backend.app.services.report_service import ReportSection, ReportTemplate


class CodeHighlightExtension(Extension):
    """Markdown扩展，用于代码高亮"""
    def extendMarkdown(self, md):
        md.registerExtension(self)
        
        # 替换代码块处理器
        md.preprocessors.register(CodeBlockPreprocessor(md), 'code_block', 25)


class CodeBlockPreprocessor:
    """代码块预处理器"""
    def __init__(self, md):
        self.md = md
        self.pattern = re.compile(r'```(\w+)?\n(.*?)\n```', re.DOTALL)
    
    def run(self, lines):
        text = '\n'.join(lines)
        
        def replace_code_block(match):
            language = match.group(1) or 'text'
            code = match.group(2)
            
            if PYGMENTS_AVAILABLE:
                try:
                    lexer = get_lexer_by_name(language, stripall=True)
                    formatter = HtmlFormatter(cssclass="codehilite")
                    result = highlight(code, lexer, formatter)
                    return result
                except Exception:
                    # 如果找不到对应的语言，回退到普通代码块
                    pass
            
            # 回退到普通代码块
            return f'<pre><code class="language-{language}">{code}</code></pre>'
        
        text = self.pattern.sub(replace_code_block, text)
        return text.split('\n')


class MarkdownFormatter:
    """Markdown格式化器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def format(self, content: str, include_toc: bool = True, include_code_highlighting: bool = True) -> str:
        """
        格式化Markdown内容
        
        Args:
            content: Markdown内容
            include_toc: 是否包含目录
            include_code_highlighting: 是否包含代码高亮
            
        Returns:
            str: 格式化后的Markdown内容
        """
        result = content
        
        # 添加目录（如果需要）
        if include_toc and "[TOC]" not in result:
            result = f"[TOC]\n\n{result}"
        
        # 添加代码高亮（如果需要）
        if include_code_highlighting:
            # 这里可以添加代码高亮的逻辑
            pass
        
        return result
    
    def generate_toc(self, content: str) -> str:
        """
        生成Markdown目录
        
        Args:
            content: Markdown内容
            
        Returns:
            str: 添加目录后的Markdown内容
        """
        # 查找所有标题
        headers = []
        for line in content.split('\n'):
            if line.startswith('#'):
                # 计算标题级别
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                # 提取标题文本
                title = line[level:].strip()
                
                # 创建锚点ID
                anchor = title.lower().replace(' ', '-').replace('.', '').replace(',', '')
                anchor = re.sub(r'[^a-z0-9-]', '', anchor)
                
                headers.append((level, title, anchor))
        
        # 生成目录
        toc = "## 目录\n\n"
        for level, title, anchor in headers:
            # 跳过目录本身
            if title == "目录":
                continue
            
            # 添加缩进
            indent = "  " * (level - 1)
            toc += f"{indent}- [{title}](#{anchor})\n"
        
        toc += "\n"
        
        # 在内容中添加锚点
        lines = content.split('\n')
        result = []
        
        for line in lines:
            if line.startswith('#'):
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                title = line[level:].strip()
                anchor = title.lower().replace(' ', '-').replace('.', '').replace(',', '')
                anchor = re.sub(r'[^a-z0-9-]', '', anchor)
                
                result.append(f'<a id="{anchor}"></a>')
            
            result.append(line)
        
        # 在第一个标题之后插入目录
        for i, line in enumerate(result):
            if line.startswith('#'):
                result.insert(i + 1, "")
                result.insert(i + 2, toc)
                break
        
        return '\n'.join(result)
    
    def highlight_code(self, content: str) -> str:
        """
        高亮Markdown代码块
        
        Args:
            content: Markdown内容
            
        Returns:
            str: 添加代码高亮后的Markdown内容
        """
        if not MARKDOWN_AVAILABLE or not PYGMENTS_AVAILABLE:
            return content
        
        # 添加CSS样式
        css = HtmlFormatter().get_style_defs('.codehilite')
        style_tag = f"<style>\n{css}\n</style>\n\n"
        
        # 转换Markdown为HTML
        extensions = ['fenced_code', CodeHighlightExtension()]
        html = markdown.markdown(content, extensions=extensions)
        
        # 返回HTML内容
        return style_tag + html
    
    def generate_table(self, headers: List[str], data: List[List[str]]) -> str:
        """
        生成Markdown表格
        
        Args:
            headers: 表头
            data: 行数据
            
        Returns:
            str: Markdown表格
        """
        # 创建表头
        header_row = "| " + " | ".join(headers) + " |"
        
        # 创建分隔行
        separator = "| " + " | ".join(["---"] * len(headers)) + " |"
        
        # 创建数据行
        data_rows = []
        for row in data:
            data_rows.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        # 组合表格
        table = header_row + "\n" + separator + "\n" + "\n".join(data_rows)
        
        return table
    
    def add_code_block(self, code: str, language: str = None) -> str:
        """
        添加代码块
        
        Args:
            code: 代码内容
            language: 代码语言
            
        Returns:
            str: 格式化后的代码块
        """
        if language:
            return f"```{language}\n{code}\n```"
        else:
            return f"```\n{code}\n```"
    
    def add_image(self, title: str, path: str, alt_text: str = "") -> str:
        """
        添加Markdown图片
        
        Args:
            title: 图片标题
            path: 图片路径
            alt_text: 替代文本
            
        Returns:
            str: Markdown图片
        """
        alt = alt_text or title
        return f"![{alt}]({path})"
    
    def add_link(self, text: str, url: str) -> str:
        """
        添加Markdown链接
        
        Args:
            text: 链接文本
            url: 链接URL
            
        Returns:
            str: Markdown链接
        """
        return f"[{text}]({url})"


class DocxFormatter:
    """DOCX格式化器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ReportGenerationException("未安装python-docx库，无法使用DOCX格式化器")
        self.logger = logging.getLogger(__name__)
    
    def format(self, doc, template, include_toc: bool = True, include_styles: bool = True) -> None:
        """
        格式化DOCX文档
        
        Args:
            doc: DOCX文档对象
            template: 报告模板
            include_toc: 是否包含目录
            include_styles: 是否包含样式
        """
        if not DOCX_AVAILABLE:
            raise ImportError("未安装python-docx库，无法格式化DOCX文档")
        
        # 添加样式（如果需要）
        if include_styles:
            self._add_styles(doc)
        
        # 添加目录（如果需要）
        if include_toc:
            self._add_toc(doc)
    
    def create_document(self, title: str) -> Any:
        """
        创建DOCX文档
        
        Args:
            title: 文档标题
            
        Returns:
            Any: DOCX文档对象
        """
        doc = docx.Document()
        doc.add_heading(title, 0)
        return doc
    
    def add_heading(self, doc, text: str, level: int = 1) -> Any:
        """
        添加标题
        
        Args:
            doc: DOCX文档对象
            text: 标题文本
            level: 标题级别
            
        Returns:
            Any: 标题对象
        """
        return doc.add_heading(text, level)
    
    def add_paragraph(self, doc, text: str) -> Any:
        """
        添加段落
        
        Args:
            doc: DOCX文档对象
            text: 段落文本
            
        Returns:
            Any: 段落对象
        """
        return doc.add_paragraph(text)
    
    def add_table(self, doc, headers: List[str], data: List[List[str]]) -> Any:
        """
        添加表格
        
        Args:
            doc: DOCX文档对象
            headers: 表头
            data: 行数据
            
        Returns:
            Any: 表格对象
        """
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # 添加行
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table
    
    def add_styles(self, doc) -> None:
        """
        添加样式
        
        Args:
            doc: DOCX文档对象
        """
        # 添加标题样式
        for i in range(1, 7):
            if f"Heading {i}" not in doc.styles:
                style = doc.styles.add_style(f"Heading {i}", WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = "Arial"
                style.font.size = Pt(20 - i * 2)  # 标题大小逐级递减
                style.font.bold = True
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)
        
        # 添加正文样式
        if "Body Text" not in doc.styles:
            style = doc.styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
        
        # 添加代码样式
        if "Code" not in doc.styles:
            style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Courier New"
            style.font.size = Pt(10)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
    
    def _add_toc(self, doc) -> None:
        """
        添加目录
        
        Args:
            doc: DOCX文档
        """
        # 添加目录标题
        doc.add_heading("目录", 1)
        
        # 添加分隔线
        doc.add_paragraph("_" * 40)
        
        # 注意：python-docx不直接支持自动生成目录
        # 这里只是添加一个占位符
        doc.add_paragraph("目录将在此处生成")
        
        # 添加分隔线
        doc.add_paragraph("_" * 40)
        doc.add_paragraph()  # 添加空行
    
    def add_code_block(self, doc, code: str, language: str = None) -> None:
        """
        添加代码块
        
        Args:
            doc: DOCX文档
            code: 代码内容
            language: 代码语言
        """
        if not DOCX_AVAILABLE:
            raise ImportError("未安装python-docx库，无法添加代码块")
        
        # 添加语言标签（如果有）
        if language:
            lang_para = doc.add_paragraph()
            lang_para.style = "Code"
            lang_run = lang_para.add_run(f"Language: {language}")
            lang_run.bold = True
        
        # 添加代码内容
        code_para = doc.add_paragraph()
        code_para.style = "Code"
        code_para.add_run(code)
    
    def add_chart(self, doc, data: Dict[str, Any], chart_type: str = "bar") -> None:
        """
        添加图表
        
        Args:
            doc: DOCX文档
            data: 图表数据
            chart_type: 图表类型
        """
        if not DOCX_AVAILABLE:
            raise ImportError("未安装python-docx库，无法添加图表")
        
        # 由于python-docx对图表支持有限，这里简单地添加图表数据作为表格
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "类别"
        header_cells[1].text = "值"
        
        # 添加数据
        for category, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(category)
            row_cells[1].text = str(value)


class ReportFormatterService:
    """报告格式化服务"""
    
    def __init__(self):
        """初始化报告格式化服务"""
        self.logger = logging.getLogger(__name__)
        self.formatters = {
            "markdown": MarkdownFormatter(),
            "docx": DocxFormatter() if DOCX_AVAILABLE else None
        }
    
    def get_formatter(self, format_type: str):
        """
        获取格式化器
        
        Args:
            format_type: 格式类型
            
        Returns:
            格式化器对象
            
        Raises:
            ValueError: 不支持的格式类型
        """
        formatter = self.formatters.get(format_type.lower())
        if not formatter:
            raise ValueError(f"不支持的格式化器类型: {format_type}")
        return formatter
    
    def format_markdown(self, content: str, include_toc: bool = True, include_code_highlighting: bool = True) -> str:
        """
        格式化Markdown内容
        
        Args:
            content: Markdown内容
            include_toc: 是否包含目录
            include_code_highlighting: 是否包含代码高亮
            
        Returns:
            str: 格式化后的Markdown内容
        """
        result = content
        
        # 添加目录（如果需要）
        if include_toc and "[TOC]" not in result:
            result = f"[TOC]\n\n{result}"
        
        # 添加代码高亮（如果需要）
        if include_code_highlighting:
            result = self._add_code_highlighting(result)
        
        return result
    
    def format_docx(self, doc, template, include_toc: bool = True, include_styles: bool = True) -> None:
        """
        格式化DOCX文档
        
        Args:
            doc: DOCX文档对象
            template: 报告模板
            include_toc: 是否包含目录
            include_styles: 是否包含样式
        """
        formatter = self.get_formatter("docx")
        formatter.format(doc, template, include_toc, include_styles)
    
    def add_table_to_markdown(self, headers: List[str], rows: List[List[str]]) -> str:
        """
        创建Markdown表格
        
        Args:
            headers: 表头
            rows: 行数据
            
        Returns:
            str: Markdown表格
        """
        formatter = self.get_formatter("markdown")
        return formatter.generate_table(headers, rows)
    
    def add_code_block_to_markdown(self, code: str, language: str = None) -> str:
        """
        创建Markdown代码块
        
        Args:
            code: 代码内容
            language: 代码语言
            
        Returns:
            str: Markdown代码块
        """
        formatter = self.get_formatter("markdown")
        return formatter.add_code_block(code, language)
    
    def markdown_to_docx(
        self, 
        markdown_content: str, 
        output_path: str,
        include_toc: bool = True,
        include_code_highlighting: bool = True,
        include_styles: bool = True,
        include_charts: bool = False,
        chart_data: Optional[Dict[str, Any]] = None
    ) -> None:
        """
        将Markdown内容转换为DOCX格式
        
        Args:
            markdown_content: Markdown内容
            output_path: 输出路径
            include_toc: 是否包含目录
            include_code_highlighting: 是否包含代码高亮
            include_styles: 是否包含样式
            include_charts: 是否包含图表
            chart_data: 图表数据
        """
        if not DOCX_AVAILABLE:
            raise ImportError("未安装python-docx库，无法转换为DOCX格式")
        
        # 创建文档
        doc = docx.Document()
        
        # 添加样式（如果需要）
        if include_styles:
            self._add_docx_styles(doc)
        
        # 解析Markdown内容
        lines = markdown_content.split("\n")
        i = 0
        
        # 处理[TOC]标记（如果有）
        if include_toc and i < len(lines) and lines[i].strip() == "[TOC]":
            i += 1  # 跳过[TOC]行
            # 这里可以添加目录生成逻辑
            # 由于python-docx不直接支持目录，这里可以手动生成一个简单的目录
            self._add_toc_to_docx(doc, markdown_content)
            
            # 跳过空行
            while i < len(lines) and not lines[i].strip():
                i += 1
        
        # 处理文档标题和正文
        current_paragraph = None
        in_code_block = False
        code_block_content = []
        code_language = None
        
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.startswith("```"):
                if not in_code_block:
                    # 开始代码块
                    in_code_block = True
                    code_block_content = []
                    # 获取代码语言
                    code_language = line[3:].strip()
                else:
                    # 结束代码块
                    in_code_block = False
                    # 添加代码块到文档
                    if include_code_highlighting:
                        self._add_code_block_to_docx(doc, code_block_content, code_language)
                    else:
                        # 简单添加为普通段落
                        code_para = doc.add_paragraph()
                        code_para.style = "Code"
                        code_para.add_run("\n".join(code_block_content))
                i += 1
                continue
            
            if in_code_block:
                # 收集代码块内容
                code_block_content.append(line)
                i += 1
                continue
            
            # 处理标题
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                doc.add_heading(title, level)
                i += 1
                continue
            
            # 处理列表项
            list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.+)$", line)
            if list_match:
                indent = len(list_match.group(1))
                list_type = list_match.group(2)
                content = list_match.group(3).strip()
                
                # 确定列表级别和类型
                level = indent // 2
                is_numbered = bool(re.match(r"\d+\.", list_type))
                
                # 添加列表项
                p = doc.add_paragraph(content)
                p.style = f"List Bullet {level + 1}" if not is_numbered else f"List Number {level + 1}"
                i += 1
                continue
            
            # 处理普通段落
            if line.strip():
                # 处理加粗、斜体等Markdown格式
                p = doc.add_paragraph()
                self._parse_markdown_formatting(p, line)
            else:
                # 空行
                doc.add_paragraph()
            
            i += 1
        
        # 添加图表（如果需要）
        if include_charts and chart_data:
            self._add_charts_to_docx(doc, chart_data)
        
        # 保存文档
        doc.save(output_path)
    
    def _add_toc_to_docx(self, doc, markdown_content: str) -> None:
        """
        向DOCX文档添加目录
        
        Args:
            doc: DOCX文档
            markdown_content: Markdown内容
        """
        # 添加目录标题
        doc.add_heading("目录", 1)
        
        # 解析Markdown中的标题
        headers = []
        for line in markdown_content.split("\n"):
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                headers.append((level, title))
        
        # 添加目录项
        for level, title in headers:
            # 跳过目录自身
            if title == "目录":
                continue
            
            # 添加目录项，根据级别缩进
            indent = "  " * (level - 1)
            p = doc.add_paragraph(f"{indent}• {title}")
            p.style = "TOC"
        
        # 添加分隔
        doc.add_paragraph()
        doc.add_paragraph("---")
        doc.add_paragraph()
    
    def _add_code_highlighting(self, content: str) -> str:
        """
        添加代码高亮
        
        Args:
            content: Markdown内容
            
        Returns:
            str: 添加代码高亮后的内容
        """
        # 这里可以添加代码高亮的逻辑
        # 例如，可以在代码块前后添加特殊标记
        return content
    
    def _add_docx_styles(self, doc) -> None:
        """
        向DOCX文档添加样式
        
        Args:
            doc: DOCX文档
        """
        # 添加代码样式
        if "Code" not in doc.styles:
            code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = "Courier New"
            code_style.font.size = Pt(10)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        
        # 添加TOC样式
        if "TOC" not in doc.styles:
            toc_style = doc.styles.add_style("TOC", WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.name = "Calibri"
            toc_style.font.size = Pt(11)
            toc_style.paragraph_format.space_before = Pt(2)
            toc_style.paragraph_format.space_after = Pt(2)
        
        # 添加列表样式
        for i in range(1, 4):
            # 项目符号列表
            if f"List Bullet {i}" not in doc.styles:
                list_style = doc.styles.add_style(f"List Bullet {i}", WD_STYLE_TYPE.PARAGRAPH)
                list_style.base_style = doc.styles["Normal"]
                list_style.paragraph_format.left_indent = Inches(i * 0.25)
                list_style.paragraph_format.first_line_indent = Inches(-0.25)
            
            # 编号列表
            if f"List Number {i}" not in doc.styles:
                list_style = doc.styles.add_style(f"List Number {i}", WD_STYLE_TYPE.PARAGRAPH)
                list_style.base_style = doc.styles["Normal"]
                list_style.paragraph_format.left_indent = Inches(i * 0.25)
                list_style.paragraph_format.first_line_indent = Inches(-0.25)
    
    def _add_code_block_to_docx(self, doc, code_lines: List[str], language: str = None) -> None:
        """
        向DOCX文档添加代码块
        
        Args:
            doc: DOCX文档
            code_lines: 代码行列表
            language: 代码语言
        """
        # 添加语言标签（如果有）
        if language:
            lang_para = doc.add_paragraph()
            lang_para.style = "Code"
            lang_run = lang_para.add_run(f"Language: {language}")
            lang_run.bold = True
        
        # 添加代码内容
        code_para = doc.add_paragraph()
        code_para.style = "Code"
        code_para.add_run("\n".join(code_lines))
    
    def _add_charts_to_docx(self, doc, chart_data: Dict[str, Any]) -> None:
        """
        向DOCX文档添加图表
        
        Args:
            doc: DOCX文档
            chart_data: 图表数据
        """
        # 添加图表标题
        doc.add_heading("图表", 1)
        
        # 添加图表
        for chart_info in chart_data.get("charts", []):
            title = chart_info.get("title", "图表")
            chart_type = chart_info.get("type", "bar")
            data = chart_info.get("data", {})
            
            # 添加图表标题
            doc.add_heading(title, 2)
            
            # 由于python-docx对图表支持有限，这里简单地添加图表数据作为表格
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            
            # 添加表头
            header_cells = table.rows[0].cells
            header_cells[0].text = "类别"
            header_cells[1].text = "值"
            
            # 添加数据
            for category, value in data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(category)
                row_cells[1].text = str(value)
            
            # 添加空行
            doc.add_paragraph()
    
    def _parse_markdown_formatting(self, paragraph, text: str) -> None:
        """
        解析Markdown格式并应用到DOCX段落
        
        Args:
            paragraph: DOCX段落
            text: Markdown文本
        """
        # 解析加粗、斜体等格式
        i = 0
        while i < len(text):
            # 检查加粗
            if i + 1 < len(text) and text[i:i+2] == "**" and "**" in text[i+2:]:
                end = text.find("**", i+2)
                if end != -1:
                    run = paragraph.add_run(text[i+2:end])
                    run.bold = True
                    i = end + 2
                    continue
            
            # 检查斜体
            if text[i] == "*" and "*" in text[i+1:]:
                end = text.find("*", i+1)
                if end != -1:
                    run = paragraph.add_run(text[i+1:end])
                    run.italic = True
                    i = end + 1
                    continue
            
            # 检查代码
            if text[i] == "`" and "`" in text[i+1:]:
                end = text.find("`", i+1)
                if end != -1:
                    run = paragraph.add_run(text[i+1:end])
                    run.font.name = "Courier New"
                    i = end + 1
                    continue
            
            # 普通文本
            j = i
            while j < len(text):
                if text[j] in ["*", "`"]:
                    break
                j += 1
            
            if j > i:
                paragraph.add_run(text[i:j])
                i = j
            else:
                paragraph.add_run(text[i])
                i += 1 

    def format_report(self, template: ReportTemplate, sections: List[ReportSection], format_type: ReportFormatEnum) -> bytes:
        """
        根据指定的格式生成报告
        """
        if format_type == ReportFormatEnum.MARKDOWN:
            return self._to_markdown(template, sections).encode('utf-8')
        elif format_type == ReportFormatEnum.DOCX:
            return self._to_docx(template, sections)
        else:
            raise ValueError(f"不支持的报告格式: {format_type}")

    def _to_markdown(self, template: ReportTemplate, sections: List[ReportSection]) -> str:
        """
        将报告转换为Markdown格式
        """
        markdown_content = ""
        
        # 使用模板的额外信息
        if template.extra_info:
            for key, value in template.extra_info.items():
                markdown_content += f"**{key.capitalize()}**: {value}\\n\\n"
        
        # 添加章节
        for section in sorted(sections, key=lambda s: s.order):
            markdown_content += f"## {section.title}\\n\\n"
            markdown_content += section.content
            markdown_content += "\\n\\n"
        
        return markdown_content

    def _to_docx(self, template: ReportTemplate, sections: List[ReportSection]) -> bytes:
        """生成DOCX格式的报告"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，无法生成DOCX报告")

        document = docx.Document()
        
        # 使用模板的额外信息
        if template.extra_info:
            for key, value in template.extra_info.items():
                document.add_paragraph(f"{key.capitalize()}: {value}")

        # 添加章节
        for section in sorted(sections, key=lambda s: s.order):
            document.add_heading(section.title, 1)
            document.add_paragraph(section.content)
        
        # 保存文档
        output_path = f"{template.title.replace(' ', '_')}.docx"
        document.save(output_path)
        
        # 读取文档内容
        with open(output_path, 'rb') as f:
            return f.read() 