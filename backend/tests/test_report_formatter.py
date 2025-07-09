import os
import pytest
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from backend.app.services.report_formatter_service import ReportFormatterService, MarkdownFormatter, DocxFormatter


class TestReportFormatter:
    """报告格式化服务测试类"""
    
    def setup_method(self):
        """每个测试方法前的设置"""
        # 创建报告格式化服务
        self.formatter_service = ReportFormatterService()
    
    def test_markdown_formatter_creation(self):
        """测试Markdown格式化器创建"""
        formatter = self.formatter_service.get_formatter("markdown")
        assert isinstance(formatter, MarkdownFormatter)
    
    def test_docx_formatter_creation(self):
        """测试DOCX格式化器创建"""
        try:
            formatter = self.formatter_service.get_formatter("docx")
            assert isinstance(formatter, DocxFormatter)
        except ImportError:
            pytest.skip("python-docx not installed")
    
    def test_invalid_formatter_type(self):
        """测试无效的格式化器类型"""
        with pytest.raises(ValueError):
            self.formatter_service.get_formatter("invalid_type")
    
    def test_markdown_toc_generation(self):
        """测试Markdown目录生成"""
        markdown_content = """
# 主标题

## 第一部分
内容1

## 第二部分
内容2

### 子部分2.1
子内容2.1

## 第三部分
内容3
"""
        formatter = self.formatter_service.get_formatter("markdown")
        toc = formatter.generate_toc(markdown_content)
        
        # 验证目录内容
        assert "## 目录" in toc
        assert "- [第一部分](#第一部分)" in toc
        assert "- [第二部分](#第二部分)" in toc
        assert "  - [子部分2.1](#子部分21)" in toc
        assert "- [第三部分](#第三部分)" in toc
    
    def test_markdown_code_highlighting(self):
        """测试Markdown代码高亮"""
        markdown_content = """
# 代码示例

```python
def hello_world():
    print("Hello, World!")
```

普通文本

```javascript
function greet() {
    console.log("Hello!");
}
```
"""
        formatter = self.formatter_service.get_formatter("markdown")
        highlighted = formatter.highlight_code(markdown_content)
        
        # 验证高亮内容
        assert "```python" not in highlighted
        assert '<pre><code class="language-python">' in highlighted
        assert "```javascript" not in highlighted
        assert '<pre><code class="language-javascript">' in highlighted
        assert "def hello_world" in highlighted
        assert "function greet" in highlighted
    
    def test_markdown_table_generation(self):
        """测试Markdown表格生成"""
        headers = ["姓名", "年龄", "职业"]
        data = [
            ["张三", "25", "工程师"],
            ["李四", "30", "设计师"],
            ["王五", "28", "产品经理"]
        ]
        
        formatter = self.formatter_service.get_formatter("markdown")
        table = formatter.generate_table(headers, data)
        
        # 验证表格内容
        assert "| 姓名 | 年龄 | 职业 |" in table
        assert "| --- | --- | --- |" in table
        assert "| 张三 | 25 | 工程师 |" in table
        assert "| 李四 | 30 | 设计师 |" in table
        assert "| 王五 | 28 | 产品经理 |" in table
    
    @pytest.mark.skipif(True, reason="需要python-docx支持")
    def test_docx_formatting(self):
        """测试DOCX格式化功能（需要python-docx支持）"""
        try:
            import docx
            
            # 创建DOCX格式化器
            formatter = self.formatter_service.get_formatter("docx")
            
            # 创建文档
            doc = formatter.create_document("测试文档")
            
            # 添加标题
            formatter.add_heading(doc, "主标题", level=1)
            formatter.add_heading(doc, "子标题", level=2)
            
            # 添加段落
            formatter.add_paragraph(doc, "这是一个测试段落。")
            
            # 添加表格
            headers = ["姓名", "年龄", "职业"]
            data = [
                ["张三", "25", "工程师"],
                ["李四", "30", "设计师"]
            ]
            formatter.add_table(doc, headers, data)
            
            # 验证文档内容
            assert len(doc.paragraphs) > 0
            assert doc.paragraphs[0].text == "主标题"
            assert doc.paragraphs[1].text == "子标题"
            assert "这是一个测试段落" in doc.paragraphs[2].text
            assert len(doc.tables) > 0
            assert len(doc.tables[0].rows) == 3  # 表头 + 2行数据
            
        except ImportError:
            pytest.skip("python-docx not installed")
    
    def test_markdown_image_insertion(self):
        """测试Markdown图片插入"""
        formatter = self.formatter_service.get_formatter("markdown")
        image_markdown = formatter.add_image("测试图片", "path/to/image.png", alt_text="测试图片")
        
        # 验证图片Markdown
        assert "![测试图片](path/to/image.png)" in image_markdown
    
    def test_markdown_link_insertion(self):
        """测试Markdown链接插入"""
        formatter = self.formatter_service.get_formatter("markdown")
        link_markdown = formatter.add_link("谷歌", "https://www.google.com")
        
        # 验证链接Markdown
        assert "[谷歌](https://www.google.com)" in link_markdown
    
    def test_code_block_generation(self):
        """测试代码块生成"""
        formatter = self.formatter_service.get_formatter("markdown")
        code = """
def hello():
    print("Hello, World!")
"""
        code_block = formatter.add_code_block(code, language="python")
        
        # 验证代码块
        assert "```python" in code_block
        assert "def hello():" in code_block
        assert "```" in code_block
    
    def test_markdown_formatting_options(self):
        """测试Markdown格式化选项"""
        content = """
# 测试文档

## 第一部分
内容1

```python
def test():
    pass
```

## 第二部分
内容2
"""
        # 测试不同选项组合
        formatter = self.formatter_service.get_formatter("markdown")
        
        # 测试只有目录
        formatted_toc_only = formatter.format(content, include_toc=True, include_code_highlighting=False)
        assert "## 目录" in formatted_toc_only
        assert "```python" in formatted_toc_only  # 代码块未高亮
        
        # 测试只有代码高亮
        formatted_highlight_only = formatter.format(content, include_toc=False, include_code_highlighting=True)
        assert "## 目录" not in formatted_highlight_only
        assert "```python" not in formatted_highlight_only
        assert '<pre><code class="language-python">' in formatted_highlight_only
        
        # 测试两者都有
        formatted_both = formatter.format(content, include_toc=True, include_code_highlighting=True)
        assert "## 目录" in formatted_both
        assert "```python" not in formatted_both
        assert '<pre><code class="language-python">' in formatted_both 